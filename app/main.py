from docx import Document

from fastapi import FastAPI, Depends
from sqlalchemy.orm import Session
from typing import List

from app.routers import admin_router, profile_router
from app.sql_app.database import get_database
from app.sql_app.releases_service import get_all_releases
from models import ReleaseStageCreate, UserCreate
from auth import get_current_user, create_user, verify_role

app = FastAPI()

app.include_router(admin_router.router)
app.include_router(profile_router.router)


#TODO:  Разнести методы по роуетрам

@app.get("/admin/approve_release/{release_id}")
def approve_release(release_id: int, db: Session = Depends(get_database), token: str = Depends(oauth2_scheme)):
    verify_role(token, "admin", db)  # Проверяем роль пользователя
    return {"message": f"Release {release_id} approved"}


# Маршрут для создания нового этапа релиза
@app.post("/release_stages/", response_model=ReleaseStageCreate)
def create_release_stage(stage: ReleaseStageCreate, db: Session = Depends(get_database),
                         current_user: str = Depends(get_current_user)):
    release = create_release_stage(stage=stage, db=db)
    return release


# Маршрут для получения всех этапов релиза
@app.get("/release_stages/", response_model=List[ReleaseStageCreate])
def get_all_release_stages(db: Session = Depends(get_database)):
    return get_all_releases(db=db)


@app.put("/release_stages/{stage_id}")
def update_release_stage(stage_id: int, name: str, description: str, start_date: str, end_date: str,
                         responsible_person: str,
                         current_user: User = Depends(get_current_active_user),
                         db: Session = Depends(get_database)):
    if current_user.role not in [RoleEnum.admin, RoleEnum.release_manager]:
        raise HTTPException(status_code=403, detail="Not enough permissions")
    stage = db.query(ReleaseStage).filter(ReleaseStage.id == stage_id).first()
    if not stage:
        raise HTTPException(status_code=404, detail="Release stage not found")

    stage.name = name
    stage.description = description
    stage.start_date = start_date
    stage.end_date = end_date
    stage.responsible_person = responsible_person
    db.commit()
    db.refresh(stage)

    # Добавляем запись в историю
    log = ReleaseLog(release_id=stage.id, user_id=current_user.id, action="updated")
    db.add(log)
    db.commit()


@app.get("/release_stages/report")
def generate_report(db: Session = Depends(get_database)):
    stages = get_all_releases(db=db)
    doc = Document()
    doc.add_heading("Release Stages Report", level=1)
    for stage in stages:
        doc.add_heading(stage.name, level=2)
        doc.add_paragraph(f"Description: {stage.description}")
        doc.add_paragraph(f"Start Date: {stage.start_date}")
        doc.add_paragraph(f"End Date: {stage.end_date}")
        doc.add_paragraph(f"Responsible Person: {stage.responsible_person}")
    file_path = "release_report.docx"
    doc.save(file_path)
    return {"message": f"Report generated at {file_path}"}
